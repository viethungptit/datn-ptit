import logging
import aiohttp
import aiofiles
import os
import tempfile
from PyPDF2 import PdfReader
from docx import Document

logger = logging.getLogger("extract_text")

async def extract_text_from_url(file_url: str) -> str:
    logger.info("Extracting text from URL: %s", file_url)
    tmp_path = await download_file(file_url)

    try:
        ext = os.path.splitext(tmp_path)[1].lower()
        logger.debug("Downloaded file saved to %s (ext=%s)", tmp_path, ext)
        if ext == ".pdf":
            text = extract_text_from_pdf(tmp_path)
        elif ext == ".docx":
            text = extract_text_from_docx(tmp_path)
        else:
            logger.warning("Unsupported file type: %s", ext)
            raise ValueError(f"Unsupported file type: {ext}")
    finally:
        try:
            os.remove(tmp_path)
            logger.debug("Temporary file removed: %s", tmp_path)
        except Exception:
            logger.exception("Failed to remove temporary file: %s", tmp_path)
    return text.strip()


async def download_file(url: str) -> str:
    filename = os.path.basename(url.split("?")[0])
    tmp_path = os.path.join(tempfile.gettempdir(), filename)
    logger.info("Downloading file from URL: %s", url)
    async with aiohttp.ClientSession() as session:
        async with session.get(url) as resp:
            if resp.status != 200:
                logger.error("Failed to download file (%s): %s", resp.status, url)
                raise Exception(f"Failed to download file ({resp.status}): {url}")
            async with aiofiles.open(tmp_path, "wb") as f:
                await f.write(await resp.read())
                logger.debug("File written to temporary path: %s", tmp_path)

    return tmp_path


def extract_text_from_pdf(path: str) -> str:
    reader = PdfReader(path)
    text = "\n".join(page.extract_text() or "" for page in reader.pages)
    logger.debug("Extracted text from PDF %s (pages=%d, chars=%d)", path, len(reader.pages), len(text))
    return text


def extract_text_from_docx(path: str) -> str:
    doc = Document(path)
    texts = []
    for p in doc.paragraphs:
        if p.text.strip():
            texts.append(p.text.strip())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    texts.append(cell.text.strip())
    seen = set()
    unique_texts = []
    for t in texts:
        if t not in seen:
            seen.add(t)
            unique_texts.append(t)
    result = "\n".join(unique_texts)
    logger.debug("Extracted text from DOCX %s (elements=%d, chars=%d)", path, len(unique_texts), len(result))
    return result
